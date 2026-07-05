"""
Minimal Word report builder for the config-driven pipeline.

This module intentionally stays thin: it creates one DOCX per SharedData row,
adds a small metadata section, selected payload fields, and the final PNG from
the image pipeline.
"""

from __future__ import annotations

from copy import deepcopy
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.table import _Row
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from models.shared_data import SharedData


INVALID_FILENAME_CHARS = re.compile(r'[<>:"/\\|?*]+')
PLACEHOLDER_PATTERN = re.compile(
    r"\{\{(?:[#/]?repeat:[A-Za-z0-9_]+|[A-Za-z0-9_]+)\}\}"
)


def build_word(
    shared_data: SharedData,
    *,
    word_config: dict,
    image_path: Path,
    base_dir: Path,
) -> Path:
    """Create a minimal DOCX report for one SharedData record."""
    image_path = Path(image_path)
    if not image_path.exists():
        raise FileNotFoundError(f"build_word input image not found: {image_path}")

    output_path = build_output_path(shared_data, word_config, base_dir)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    if word_config.get("template_path"):
        document = build_document_from_template(
            shared_data,
            word_config=word_config,
            image_path=image_path,
            base_dir=base_dir,
        )
    else:
        document = build_document_from_scratch(
            shared_data,
            word_config=word_config,
            image_path=image_path,
            base_dir=base_dir,
        )

    document.save(output_path)
    return output_path


def build_document_from_template(
    shared_data: SharedData,
    *,
    word_config: dict,
    image_path: Path,
    base_dir: Path,
) -> Document:
    """Build a DOCX by replacing placeholders in a configured template."""
    template_path = resolve_project_path(base_dir, str(word_config["template_path"]))
    if not template_path.exists():
        raise FileNotFoundError(f"word template not found: {template_path}")

    document = Document(template_path)
    expand_repeat_blocks(document, shared_data, word_config)
    replacements = build_placeholder_replacements(shared_data, word_config)
    replace_text_placeholders(document, replacements)
    replace_image_placeholders(document, image_path, word_config)
    return document


def build_document_from_scratch(
    shared_data: SharedData,
    *,
    word_config: dict,
    image_path: Path,
    base_dir: Path,
) -> Document:
    """Build the original MVP DOCX without a template."""
    document = Document()
    configure_document(document)

    add_title(document, word_config)
    add_metadata_table(document, shared_data)
    add_payload_table(document, shared_data, word_config)
    add_image(document, image_path, word_config)
    return document


def configure_document(document: Document) -> None:
    """Apply simple page and type defaults."""
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)


def add_title(document: Document, word_config: dict) -> None:
    title = str(word_config.get("title", "MVP Image Report")).strip()
    paragraph = document.add_heading(title, level=0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_metadata_table(document: Document, shared_data: SharedData) -> None:
    document.add_heading("Metadata", level=1)
    rows = [
        ("record_id", shared_data.record_id),
        ("project_id", shared_data.project_id),
        ("test_id", shared_data.test_id),
        ("batch_sequence_id", shared_data.batch_sequence_id),
    ]
    add_key_value_table(document, rows)


def add_payload_table(
    document: Document,
    shared_data: SharedData,
    word_config: dict,
) -> None:
    fields = word_config.get("fields", [])
    if not isinstance(fields, list) or not fields:
        return

    document.add_heading("Payload", level=1)
    rows: list[tuple[str, Any]] = []
    for field_config in fields:
        if not isinstance(field_config, dict):
            continue

        source = str(field_config.get("source", "")).strip()
        if not source:
            continue

        label = str(field_config.get("label", source)).strip() or source
        value = shared_data.get_value(source)
        rows.append((label, "" if value is None else value))

    if rows:
        add_key_value_table(document, rows)


def add_key_value_table(document: Document, rows: list[tuple[str, Any]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(label)
        cells[1].text = str(value)
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_paragraph()


def add_image(document: Document, image_path: Path, word_config: dict) -> None:
    image_config = word_config.get("image", {})
    if not isinstance(image_config, dict):
        image_config = {}

    caption = str(image_config.get("caption", "Generated image")).strip()
    if caption:
        paragraph = document.add_heading(caption, level=1)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    width_inches = float(image_config.get("width_inches", 6.5))
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = image_paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def build_placeholder_replacements(
    shared_data: SharedData,
    word_config: dict,
) -> dict[str, str]:
    placeholders = word_config.get("placeholders", {})
    if not isinstance(placeholders, dict):
        return {}

    replacements: dict[str, str] = {}
    for placeholder, rule in placeholders.items():
        if not isinstance(rule, dict):
            continue

        replacements[str(placeholder)] = resolve_placeholder_value(
            shared_data,
            word_config,
            rule,
        )

    return replacements


def resolve_placeholder_value(
    shared_data: SharedData,
    word_config: dict,
    rule: dict,
) -> str:
    if "value" in rule:
        return "" if rule["value"] is None else str(rule["value"])

    source = str(rule.get("source", "")).strip()
    if not source:
        return ""

    built_in_values = {
        "record_id": shared_data.record_id,
        "project_id": shared_data.project_id,
        "test_id": shared_data.test_id,
        "batch_sequence_id": shared_data.batch_sequence_id,
        "title": word_config.get("title", ""),
    }
    if source in built_in_values:
        value = built_in_values[source]
    else:
        value = shared_data.get_value(source)

    return "" if value is None else str(value)


def expand_repeat_blocks(
    document: Document,
    shared_data: SharedData,
    word_config: dict,
) -> None:
    repeat_rules = word_config.get("repeat_block_rules")
    if repeat_rules is None:
        return

    if not isinstance(repeat_rules, list):
        raise ValueError("repeat_block_rules must be a list")

    for rule in repeat_rules:
        if not isinstance(rule, dict):
            raise ValueError("repeat_block_rules entries must be objects")

        expand_repeat_block(document, shared_data, rule)


def expand_repeat_block(
    document: Document,
    shared_data: SharedData,
    rule: dict,
) -> None:
    block_key = str(rule.get("block_key") or rule.get("section_key") or "").strip()
    start_marker = str(rule.get("start_marker", "")).strip()
    end_marker = str(rule.get("end_marker", "")).strip()
    required = bool(rule.get("required", False))

    if not start_marker or not end_marker:
        if required:
            raise ValueError(f"repeat block {block_key or '<unknown>'} missing markers")
        return

    block_rows = find_repeat_block_rows(document, start_marker, end_marker, block_key)
    if block_rows is None:
        if required:
            raise ValueError(
                f"repeat block {block_key or start_marker} start marker not found: "
                f"{start_marker}"
            )
        return

    table, start_index, end_index = block_rows
    rows = list(table.rows)
    template_rows = rows[start_index + 1 : end_index]
    if not template_rows:
        raise ValueError(f"repeat block {block_key or start_marker} has no template rows")

    repeat_items = resolve_repeat_items(shared_data, rule, block_key)
    template_trs = [deepcopy(row._tr) for row in template_rows]
    original_trs = [row._tr for row in rows[start_index : end_index + 1]]
    insertion_anchor = rows[start_index]._tr

    try:
        for item in repeat_items:
            replacements = build_repeat_placeholder_replacements(item, rule, block_key)
            for template_tr in template_trs:
                cloned_tr = deepcopy(template_tr)
                insertion_anchor.addprevious(cloned_tr)
                cloned_row = _Row(cloned_tr, table)
                replace_placeholders_in_table_row(cloned_row, replacements, block_key)

        for original_tr in original_trs:
            parent = original_tr.getparent()
            if parent is not None:
                parent.remove(original_tr)
    except Exception as exc:
        raise ValueError(f"repeat block {block_key or start_marker} expansion failed: {exc}") from exc

    remove_repeat_marker_paragraphs(document, {start_marker, end_marker})


def find_repeat_block_rows(
    document: Document,
    start_marker: str,
    end_marker: str,
    block_key: str,
):
    start_matches = find_marker_rows(document, start_marker)
    end_matches = find_marker_rows(document, end_marker)

    if not start_matches:
        return None

    if len(start_matches) > 1:
        raise ValueError(f"repeat block {block_key or start_marker} has multiple start markers")

    if not end_matches:
        raise ValueError(f"repeat block {block_key or start_marker} end marker not found: {end_marker}")

    if len(end_matches) > 1:
        raise ValueError(f"repeat block {block_key or start_marker} has multiple end markers")

    start_table, start_index = start_matches[0]
    end_table, end_index = end_matches[0]
    if start_table._tbl is not end_table._tbl:
        raise ValueError(f"repeat block {block_key or start_marker} markers are not in the same table")

    if start_index >= end_index:
        raise ValueError(f"repeat block {block_key or start_marker} start marker must be before end marker")

    return start_table, start_index, end_index


def find_marker_rows(document: Document, marker: str):
    matches = []
    for table in document.tables:
        for row_index, row in enumerate(table.rows):
            if marker in table_row_text(row):
                matches.append((table, row_index))

    return matches


def table_row_text(row) -> str:
    return "\n".join(cell.text for cell in row.cells)


def resolve_repeat_items(
    shared_data: SharedData,
    rule: dict,
    block_key: str,
) -> list[dict[str, Any]]:
    repeat_source = str(rule.get("repeat_source", "")).strip()
    if not repeat_source:
        raise ValueError(f"repeat block {block_key or '<unknown>'} missing repeat_source")

    value = resolve_shared_data_source(shared_data, repeat_source)
    if not isinstance(value, list):
        raise ValueError(f"repeat block {block_key or repeat_source} repeat_source must be a list")

    if not value:
        raise ValueError(f"repeat block {block_key or repeat_source} repeat_source is empty")

    repeat_items: list[dict[str, Any]] = []
    for index, item in enumerate(value, start=1):
        if not isinstance(item, dict):
            raise ValueError(
                f"repeat block {block_key or repeat_source} item {index} must be an object"
            )
        repeat_items.append(item)

    return repeat_items


def resolve_shared_data_source(shared_data: SharedData, source: str) -> Any:
    normalized = source.strip()
    if normalized.startswith("shared_data."):
        normalized = normalized.removeprefix("shared_data.")

    built_in_values = {
        "record_id": shared_data.record_id,
        "project_id": shared_data.project_id,
        "test_id": shared_data.test_id,
        "batch_sequence_id": shared_data.batch_sequence_id,
    }
    if normalized in built_in_values:
        return built_in_values[normalized]

    return resolve_nested_value(shared_data.payload, normalized)


def build_repeat_placeholder_replacements(
    repeat_item: dict[str, Any],
    rule: dict,
    block_key: str,
) -> dict[str, str]:
    replacements = {
        f"{{{{{key}}}}}": "" if value is None else str(value)
        for key, value in repeat_item.items()
        if isinstance(key, str)
    }

    placeholder_rules = rule.get("placeholder_rules", [])
    if not isinstance(placeholder_rules, list):
        raise ValueError(f"repeat block {block_key or '<unknown>'} placeholder_rules must be a list")

    for placeholder_rule in placeholder_rules:
        if not isinstance(placeholder_rule, dict):
            raise ValueError(
                f"repeat block {block_key or '<unknown>'} placeholder_rules entries must be objects"
            )

        placeholder = str(placeholder_rule.get("placeholder", "")).strip()
        if not placeholder:
            continue

        source = str(placeholder_rule.get("source", "")).strip()
        value = resolve_repeat_item_source(repeat_item, source)
        if value is None and bool(placeholder_rule.get("required", False)):
            raise ValueError(
                f"repeat block {block_key or '<unknown>'} missing required value for {placeholder}"
            )

        replacements[placeholder] = "" if value is None else str(value)

    return replacements


def resolve_repeat_item_source(repeat_item: dict[str, Any], source: str) -> Any:
    normalized = source.strip()
    if "[]" in normalized:
        normalized = normalized.split("[]", 1)[1].lstrip(".")

    return resolve_nested_value(repeat_item, normalized)


def resolve_nested_value(data: Any, path: str) -> Any:
    current = data
    for part in path.split("."):
        if not part:
            continue

        if isinstance(current, dict):
            current = current.get(part)
        else:
            return None

        if current is None:
            return None

    return current


def replace_placeholders_in_table_row(
    row,
    replacements: dict[str, str],
    block_key: str,
) -> None:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            replace_paragraph_text_placeholders(paragraph, replacements)

    unresolved = set(PLACEHOLDER_PATTERN.findall(table_row_text(row)))
    if unresolved:
        unresolved_text = ", ".join(sorted(unresolved))
        raise ValueError(
            f"repeat block {block_key or '<unknown>'} has unresolved row placeholders: "
            f"{unresolved_text}"
        )


def remove_repeat_marker_paragraphs(document: Document, markers: set[str]) -> None:
    for paragraph in list(iter_text_placeholder_paragraphs(document)):
        if any(marker in paragraph.text for marker in markers):
            remove_paragraph(paragraph)


def remove_paragraph(paragraph) -> None:
    parent = paragraph._element.getparent()
    if parent is not None:
        parent.remove(paragraph._element)


def replace_text_placeholders(document: Document, replacements: dict[str, str]) -> None:
    if not replacements:
        return

    for paragraph in iter_text_placeholder_paragraphs(document):
        replace_paragraph_text_placeholders(paragraph, replacements)


def replace_paragraph_text_placeholders(paragraph, replacements: dict[str, str]) -> None:
    text = paragraph.text
    replaced_text = text
    for placeholder, value in replacements.items():
        replaced_text = replaced_text.replace(placeholder, value)

    if replaced_text != text:
        set_paragraph_text(paragraph, replaced_text)


def replace_image_placeholders(
    document: Document,
    image_path: Path,
    word_config: dict,
) -> None:
    image_placeholders = word_config.get("image_placeholders", {})
    if not isinstance(image_placeholders, dict):
        return

    fallback_image_config = word_config.get("image", {})
    fallback_width = 6.5
    if isinstance(fallback_image_config, dict):
        fallback_width = float(fallback_image_config.get("width_inches", fallback_width))

    for paragraph in iter_paragraphs(document):
        for placeholder, rule in image_placeholders.items():
            placeholder_text = str(placeholder)
            if placeholder_text not in paragraph.text:
                continue

            width_inches = fallback_width
            if isinstance(rule, dict):
                width_inches = float(rule.get("width_inches", fallback_width))

            replace_paragraph_image_placeholder(
                paragraph,
                placeholder_text,
                image_path,
                width_inches,
            )


def replace_paragraph_image_placeholder(
    paragraph,
    placeholder: str,
    image_path: Path,
    width_inches: float,
) -> None:
    before, _, after = paragraph.text.partition(placeholder)
    paragraph.clear()
    if before:
        paragraph.add_run(before)

    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    if after:
        paragraph.add_run(after)


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def iter_paragraphs(document: Document):
    yield from iter_container_paragraphs(document)


def iter_text_placeholder_paragraphs(document: Document):
    yield from iter_paragraphs(document)

    for section in document.sections:
        for story_name in (
            "header",
            "footer",
            "first_page_header",
            "first_page_footer",
            "even_page_header",
            "even_page_footer",
        ):
            yield from iter_container_paragraphs(getattr(section, story_name))


def iter_container_paragraphs(container):
    for paragraph in container.paragraphs:
        yield paragraph

    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def build_output_path(
    shared_data: SharedData,
    word_config: dict,
    base_dir: Path,
) -> Path:
    output_dir = resolve_project_path(
        base_dir,
        str(word_config.get("output_dir", "output/pipeline_mvp/word")),
    )
    filename_pattern = str(word_config.get("filename_pattern", "{record_id}.docx"))
    filename = render_filename_pattern(filename_pattern, shared_data)

    if not filename.strip():
        filename = shared_data.record_id

    if not filename.lower().endswith(".docx"):
        filename = f"{filename}.docx"

    return output_dir / sanitize_filename(filename)


def render_filename_pattern(pattern: str, shared_data: SharedData) -> str:
    values = {
        "record_id": shared_data.record_id,
        "project_id": shared_data.project_id,
        "test_id": shared_data.test_id,
        "batch_sequence_id": shared_data.batch_sequence_id,
    }
    values.update(
        {
            key: "" if value is None else str(value)
            for key, value in shared_data.payload.items()
        }
    )
    return pattern.format_map(DefaultFormatValues(values))


def sanitize_filename(filename: str) -> str:
    filename = INVALID_FILENAME_CHARS.sub("_", filename).strip()
    return filename or "report.docx"


def resolve_project_path(base_dir: Path, reference: str) -> Path:
    path = Path(reference)
    if path.is_absolute():
        return path

    return (Path(base_dir) / path).resolve()


class DefaultFormatValues(dict):
    def __missing__(self, key: str) -> str:
        return ""
