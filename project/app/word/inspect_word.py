"""
Minimal DOCX inspection tool for the Word MVP output.

This intentionally performs structural checks only. It does not render Word
pages or require LibreOffice.
"""

from __future__ import annotations

import argparse
import sys
import zipfile
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document


DEFAULT_OUTPUT_DIR = Path(__file__).resolve().parents[2] / "output" / "pipeline_mvp" / "word"
DEFAULT_REQUIRED_TEXT = (
    "Cycle Diagram Report",
    "record_id",
    "project_id",
    "test_id",
    "batch_sequence_id",
    "Cycle Count",
)


@dataclass
class InspectionResult:
    path: Path
    failures: list[str] = field(default_factory=list)

    @property
    def passed(self) -> bool:
        return not self.failures


def main(argv: list[str] | None = None) -> int:
    args = parse_args(argv)
    output_dir = Path(args.output_dir)
    required_text = tuple(args.required_text or DEFAULT_REQUIRED_TEXT)

    results = inspect_output_dir(output_dir, required_text)
    print_results(results)

    return 0 if results and all(result.passed for result in results) else 1


def parse_args(argv: list[str] | None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Inspect generated Word MVP DOCX files."
    )
    parser.add_argument(
        "--output-dir",
        default=str(DEFAULT_OUTPUT_DIR),
        help="Directory containing generated .docx files.",
    )
    parser.add_argument(
        "--required-text",
        action="append",
        help=(
            "Required text fragment. Can be passed multiple times. "
            "Defaults to the MVP title and key field labels."
        ),
    )
    return parser.parse_args(argv)


def inspect_output_dir(
    output_dir: Path,
    required_text: tuple[str, ...],
) -> list[InspectionResult]:
    if not output_dir.exists():
        return [
            InspectionResult(
                path=output_dir,
                failures=[f"output directory does not exist: {output_dir}"],
            )
        ]

    docx_paths = sorted(output_dir.glob("*.docx"))
    if not docx_paths:
        return [
            InspectionResult(
                path=output_dir,
                failures=[f"no .docx files found in: {output_dir}"],
            )
        ]

    return [inspect_docx(path, required_text) for path in docx_paths]


def inspect_docx(path: Path, required_text: tuple[str, ...]) -> InspectionResult:
    result = InspectionResult(path=path)

    if not path.exists():
        result.failures.append("file does not exist")
        return result

    try:
        document = Document(path)
    except Exception as exc:
        result.failures.append(f"python-docx could not open file: {exc}")
        return result

    document_text = extract_document_text(document)
    for text_fragment in required_text:
        if text_fragment not in document_text:
            result.failures.append(f"missing required text: {text_fragment}")

    if not has_embedded_image(path):
        result.failures.append("no embedded image found under word/media/")

    return result


def extract_document_text(document: Document) -> str:
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text:
            text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    text_parts.append(cell.text)

    return "\n".join(text_parts)


def has_embedded_image(path: Path) -> bool:
    try:
        with zipfile.ZipFile(path) as docx_zip:
            return any(
                name.startswith("word/media/") and not name.endswith("/")
                for name in docx_zip.namelist()
            )
    except zipfile.BadZipFile:
        return False


def print_results(results: list[InspectionResult]) -> None:
    for result in results:
        if result.passed:
            print(f"PASS {result.path}")
            continue

        print(f"FAIL {result.path}")
        for failure in result.failures:
            print(f"  - {failure}")

    passed_count = sum(1 for result in results if result.passed)
    print(f"Summary: {passed_count}/{len(results)} passed")


if __name__ == "__main__":
    sys.exit(main())
