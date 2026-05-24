"""
Create the minimal DOCX template used by the Word MVP.

This script is intentionally small and deterministic so the template can be
regenerated when the placeholder contract changes.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DEFAULT_TEMPLATE_PATH = (
    Path(__file__).resolve().parents[2] / "templates" / "word" / "mvp_report_template.docx"
)


def create_template(output_path: Path = DEFAULT_TEMPLATE_PATH) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_document(document)

    title = document.add_heading("{{report_title}}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_heading("Metadata", level=1)
    metadata_table = document.add_table(rows=4, cols=2)
    metadata_table.style = "Table Grid"
    metadata_rows = [
        ("record_id", "{{record_id}}"),
        ("project_id", "{{project_id}}"),
        ("test_id", "{{test_id}}"),
        ("batch_sequence_id", "{{batch_sequence_id}}"),
    ]
    fill_table(metadata_table, metadata_rows)

    document.add_heading("Payload", level=1)
    payload_table = document.add_table(rows=1, cols=2)
    payload_table.style = "Table Grid"
    fill_table(payload_table, [("Cycle Count", "{{cycle_count}}")])

    document.add_heading("Generated Cycle Diagram", level=1)
    image_placeholder = document.add_paragraph("{{cycle_image}}")
    image_placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.save(output_path)
    return output_path


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)


def fill_table(table, rows: list[tuple[str, str]]) -> None:
    for row, (label, placeholder) in zip(table.rows, rows, strict=True):
        row.cells[0].text = label
        row.cells[1].text = placeholder


if __name__ == "__main__":
    path = create_template()
    print(f"Created template: {path}")
